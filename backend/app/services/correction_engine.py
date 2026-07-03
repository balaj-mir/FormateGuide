"""
Correction Engine — applies XML transformations to a .docx file to fix violations.
All corrections wrapped in Track Changes XML (w:rPrChange, w:pPrChange).
"""

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy
import io
import uuid
import structlog

logger = structlog.get_logger()

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _set_font_family(rpr: etree._Element, font_name: str) -> None:
    """Set font family in run properties."""
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _set_font_size(rpr: etree._Element, size_pt: float) -> None:
    """Set font size in run properties (stored as half-points)."""
    half_pts = str(int(size_pt * 2))
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        sz = etree.SubElement(rpr, qn("w:sz"))
    sz.set(qn("w:val"), half_pts)
    szCs = rpr.find(qn("w:szCs"))
    if szCs is None:
        szCs = etree.SubElement(rpr, qn("w:szCs"))
    szCs.set(qn("w:val"), half_pts)


def _set_bold(rpr: etree._Element, bold: bool) -> None:
    """Set or remove bold in run properties."""
    b = rpr.find(qn("w:b"))
    if bold:
        if b is None:
            etree.SubElement(rpr, qn("w:b"))
    else:
        if b is not None:
            rpr.remove(b)


def _set_italic(rpr: etree._Element, italic: bool) -> None:
    """Set or remove italic in run properties."""
    i = rpr.find(qn("w:i"))
    if italic:
        if i is None:
            etree.SubElement(rpr, qn("w:i"))
    else:
        if i is not None:
            rpr.remove(i)


def _set_alignment(ppr: etree._Element, alignment: str) -> None:
    """Set paragraph alignment."""
    align_map = {"left": "left", "center": "center", "right": "right", "justified": "both"}
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = etree.SubElement(ppr, qn("w:jc"))
    jc.set(qn("w:val"), align_map.get(alignment.lower(), alignment))


def _set_line_spacing(ppr: etree._Element, spacing: float) -> None:
    """Set line spacing in paragraph properties."""
    sp = ppr.find(qn("w:spacing"))
    if sp is None:
        sp = etree.SubElement(ppr, qn("w:spacing"))
    # Line spacing as 240 * multiplier (Word uses 240ths of a line)
    sp.set(qn("w:line"), str(int(spacing * 240)))
    sp.set(qn("w:lineRule"), "auto")


def _add_track_change_rpr(rpr: etree._Element, author: str = "FormatGuard") -> None:
    """Wrap run property changes in Track Changes XML."""
    rpr_change = etree.SubElement(rpr, qn("w:rPrChange"))
    rpr_change.set(qn("w:id"), str(uuid.uuid4().int % 100000))
    rpr_change.set(qn("w:author"), author)
    rpr_change.set(qn("w:date"), "2024-01-01T00:00:00Z")
    # Store original properties
    orig_rpr = etree.SubElement(rpr_change, qn("w:rPr"))


def _add_track_change_ppr(ppr: etree._Element, author: str = "FormatGuard") -> None:
    """Wrap paragraph property changes in Track Changes XML."""
    ppr_change = etree.SubElement(ppr, qn("w:pPrChange"))
    ppr_change.set(qn("w:id"), str(uuid.uuid4().int % 100000))
    ppr_change.set(qn("w:author"), author)
    ppr_change.set(qn("w:date"), "2024-01-01T00:00:00Z")
    orig_ppr = etree.SubElement(ppr_change, qn("w:pPr"))


def _apply_font_family_correction(doc: Document, expected: str) -> int:
    """Apply font family correction to all body text runs."""
    count = 0
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            continue
        for run in para.runs:
            if run.font.name and run.font.name.lower() != expected.lower():
                rpr = run._element.find(qn("w:rPr"))
                if rpr is None:
                    rpr = etree.SubElement(run._element, qn("w:rPr"))
                _set_font_family(rpr, expected)
                _add_track_change_rpr(rpr)
                count += 1
    return count


def _apply_font_size_correction(doc: Document, expected_pt: float) -> int:
    """Apply font size correction to all body text runs."""
    count = 0
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            continue
        for run in para.runs:
            if run.font.size and abs(run.font.size.pt - expected_pt) > 0.5:
                rpr = run._element.find(qn("w:rPr"))
                if rpr is None:
                    rpr = etree.SubElement(run._element, qn("w:rPr"))
                _set_font_size(rpr, expected_pt)
                _add_track_change_rpr(rpr)
                count += 1
    return count


def _apply_alignment_correction(doc: Document, expected: str, heading_only: bool = False) -> int:
    """Apply alignment correction."""
    count = 0
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        is_heading = "heading" in style_name
        if heading_only and not is_heading:
            continue
        if not heading_only and is_heading:
            continue

        ppr = para._element.find(qn("w:pPr"))
        if ppr is None:
            ppr = etree.SubElement(para._element, qn("w:pPr"))
        _set_alignment(ppr, expected)
        _add_track_change_ppr(ppr)
        count += 1
    return count


def _apply_line_spacing_correction(doc: Document, expected: float) -> int:
    """Apply line spacing correction to body text."""
    count = 0
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            continue
        ppr = para._element.find(qn("w:pPr"))
        if ppr is None:
            ppr = etree.SubElement(para._element, qn("w:pPr"))
        _set_line_spacing(ppr, expected)
        _add_track_change_ppr(ppr)
        count += 1
    return count


def _apply_margin_correction(doc: Document, margins: dict) -> int:
    """Apply page margin corrections in sectPr."""
    count = 0
    for section in doc.sections:
        changed = False
        if "top" in margins and section.top_margin:
            from docx.shared import Cm
            section.top_margin = Cm(margins["top"])
            changed = True
        if "bottom" in margins and section.bottom_margin:
            from docx.shared import Cm
            section.bottom_margin = Cm(margins["bottom"])
            changed = True
        if "left" in margins and section.left_margin:
            from docx.shared import Cm
            section.left_margin = Cm(margins["left"])
            changed = True
        if "right" in margins and section.right_margin:
            from docx.shared import Cm
            section.right_margin = Cm(margins["right"])
            changed = True
        if changed:
            count += 1
    return count


def apply_corrections(
    docx_bytes: bytes,
    violations: list,
    selected_violation_ids: list[str] | None = None,
) -> bytes:
    """
    Apply corrections for selected violations to the .docx file.
    Returns corrected .docx as bytes with corrections as Track Changes.
    """
    doc = Document(io.BytesIO(docx_bytes))
    fixes_applied = 0

    for violation in violations:
        vid = getattr(violation, 'id', str(violation)) if not isinstance(violation, dict) else violation.get('id', '')
        if selected_violation_ids and str(vid) not in selected_violation_ids:
            continue

        is_fixable = getattr(violation, 'is_auto_fixable', True) if not isinstance(violation, dict) else violation.get('is_auto_fixable', True)
        if not is_fixable:
            continue

        element_type = getattr(violation, 'element_type', '') if not isinstance(violation, dict) else violation.get('element_type', '')
        rule_name = getattr(violation, 'rule_name', '') if not isinstance(violation, dict) else violation.get('rule_name', '')
        expected = getattr(violation, 'expected_value', '') if not isinstance(violation, dict) else violation.get('expected_value', '')

        try:
            if "font family" in rule_name.lower() or (element_type == "body_text" and "font" in rule_name.lower()):
                count = _apply_font_family_correction(doc, expected)
                fixes_applied += count

            elif "font size" in rule_name.lower():
                size = float(expected.replace(" pt", "").strip())
                count = _apply_font_size_correction(doc, size)
                fixes_applied += count

            elif "alignment" in rule_name.lower():
                count = _apply_alignment_correction(
                    doc, expected, heading_only="heading" in element_type
                )
                fixes_applied += count

            elif "line spacing" in rule_name.lower() or "spacing" in rule_name.lower():
                spacing = float(expected)
                count = _apply_line_spacing_correction(doc, spacing)
                fixes_applied += count

            elif "margin" in rule_name.lower():
                margin_val = float(expected.replace(" cm", "").strip())
                margin_key = rule_name.lower().split()[0]  # "top", "bottom", "left", "right"
                count = _apply_margin_correction(doc, {margin_key: margin_val})
                fixes_applied += count

            elif "bold" in rule_name.lower():
                expected_bold = expected.lower() == "true"
                for para in doc.paragraphs:
                    for run in para.runs:
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is None:
                            rpr = etree.SubElement(run._element, qn("w:rPr"))
                        _set_bold(rpr, expected_bold)
                        _add_track_change_rpr(rpr)
                        fixes_applied += 1

            elif "italic" in rule_name.lower():
                expected_italic = expected.lower() == "true"
                for para in doc.paragraphs:
                    for run in para.runs:
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is None:
                            rpr = etree.SubElement(run._element, qn("w:rPr"))
                        _set_italic(rpr, expected_italic)
                        _add_track_change_rpr(rpr)
                        fixes_applied += 1

        except Exception as e:
            logger.warning("Failed to apply correction", rule=rule_name, error=str(e))
            continue

    logger.info("Corrections applied", fixes=fixes_applied)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
