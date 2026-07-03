"""
Document Parser — converts a .docx file into a StructuredDocumentObject
containing all formatting metadata needed for compliance checking.

Pipeline:
  1. Validate file (MIME type, ZIP structure)
  2. Decompress ZIP archive
  3. Parse word/document.xml with lxml
  4. For each paragraph: extract w:pPr + w:rPr, resolve style inheritance
  5. Parse word/styles.xml — build full style inheritance chain
  6. Parse word/sectPr — page dimensions, margins, column layout
  7. Parse all header/footer XML files
  8. Extract figure captions, table captions, reference list entries
  9. Return StructuredDocumentObject dataclass
"""

from dataclasses import dataclass, field
from typing import Optional
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import zipfile
import io
import re
import structlog

logger = structlog.get_logger()

# Word XML namespaces
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_MAP = {"w": WORD_NS}

# EMU to cm conversion (1 cm = 360000 EMU)
EMU_PER_CM = 360000
# Twips to cm (1 cm = 567 twips)
TWIPS_PER_CM = 567
# Half-points to points
HALF_PT = 2


@dataclass
class ParagraphFormatting:
    paragraph_index: int
    page_number: Optional[int]
    text_preview: str
    style_name: Optional[str]
    font_family: Optional[str]
    font_size_pt: Optional[float]
    bold: Optional[bool]
    italic: Optional[bool]
    underline: Optional[bool]
    color: Optional[str]
    alignment: Optional[str]
    line_spacing: Optional[float]
    space_before_pt: Optional[float]
    space_after_pt: Optional[float]
    indent_left_cm: Optional[float]
    indent_right_cm: Optional[float]
    indent_first_line_cm: Optional[float]
    numbering_id: Optional[int]
    numbering_level: Optional[int]
    is_heading: bool = False
    heading_level: Optional[int] = None
    is_caption: bool = False
    is_reference: bool = False
    is_list_item: bool = False


@dataclass
class PageSetup:
    width_cm: float
    height_cm: float
    margin_top_cm: float
    margin_bottom_cm: float
    margin_left_cm: float
    margin_right_cm: float
    orientation: str
    paper_size: str


@dataclass
class HeaderFooterContent:
    has_header: bool
    has_footer: bool
    header_text: Optional[str]
    footer_text: Optional[str]
    page_number_position: Optional[str]
    page_number_format: Optional[str]


@dataclass
class StructuredDocumentObject:
    filename: str
    page_count: int
    word_count: int
    paragraph_count: int
    page_setup: PageSetup
    header_footer: HeaderFooterContent
    paragraphs: list[ParagraphFormatting] = field(default_factory=list)
    has_table_of_contents: bool = False
    has_references_section: bool = False
    detected_sections: list[str] = field(default_factory=list)
    raw_xml: Optional[str] = None


def _validate_docx(file_bytes: bytes) -> None:
    """Validate that the file is a valid .docx (ZIP with proper structure)."""
    try:
        bio = io.BytesIO(file_bytes)
        if not zipfile.is_zipfile(bio):
            raise ValueError("File is not a valid ZIP archive")
        with zipfile.ZipFile(bio, "r") as zf:
            names = zf.namelist()
            if "word/document.xml" not in names:
                raise ValueError("File is not a valid .docx — missing word/document.xml")
    except zipfile.BadZipFile:
        raise ValueError("Corrupted ZIP file — cannot open as .docx")


def _emu_to_cm(emu: int) -> float:
    """Convert EMU (English Metric Units) to centimeters."""
    return round(emu / EMU_PER_CM, 2)


def _twips_to_cm(twips: int) -> float:
    """Convert twips to centimeters."""
    return round(twips / TWIPS_PER_CM, 2)


def _twips_to_pt(twips: int) -> float:
    """Convert twips to points (1 pt = 20 twips)."""
    return round(twips / 20, 1)


def _half_pt_to_pt(half_pt: int) -> float:
    """Convert half-points to points."""
    return round(half_pt / 2, 1)


def _detect_paper_size(width_cm: float, height_cm: float) -> str:
    """Detect paper size from dimensions."""
    if abs(width_cm - 21.0) < 0.5 and abs(height_cm - 29.7) < 0.5:
        return "A4"
    if abs(width_cm - 21.59) < 0.5 and abs(height_cm - 27.94) < 0.5:
        return "Letter"
    if abs(width_cm - 21.59) < 0.5 and abs(height_cm - 35.56) < 0.5:
        return "Legal"
    return "Unknown"


def _build_style_map(doc: Document) -> dict:
    """Build a map of style_id → style properties for inheritance resolution."""
    style_map = {}
    for style in doc.styles:
        try:
            sid = style.style_id
            props = {
                "name": style.name,
                "font_family": None,
                "font_size_pt": None,
                "bold": None,
                "italic": None,
                "alignment": None,
                "line_spacing": None,
                "base_style_id": style.base_style.style_id if style.base_style else None,
            }
            if style.font:
                if style.font.name:
                    props["font_family"] = style.font.name
                if style.font.size:
                    props["font_size_pt"] = style.font.size.pt
                if style.font.bold is not None:
                    props["bold"] = style.font.bold
                if style.font.italic is not None:
                    props["italic"] = style.font.italic
            if hasattr(style, "paragraph_format") and style.paragraph_format:
                pf = style.paragraph_format
                if pf.alignment is not None:
                    align_map = {0: "left", 1: "center", 2: "right", 3: "justified"}
                    props["alignment"] = align_map.get(pf.alignment, str(pf.alignment))
                if pf.line_spacing is not None:
                    props["line_spacing"] = float(pf.line_spacing)
            style_map[sid] = props
        except Exception:
            continue
    return style_map


def _resolve_property(style_map: dict, style_id: str, prop: str, max_depth: int = 10) -> any:
    """Resolve a property through the style inheritance chain."""
    visited = set()
    current = style_id
    depth = 0
    while current and depth < max_depth and current not in visited:
        visited.add(current)
        style = style_map.get(current)
        if not style:
            break
        if style.get(prop) is not None:
            return style[prop]
        current = style.get("base_style_id")
        depth += 1
    return None


def _get_alignment_str(alignment) -> Optional[str]:
    """Convert python-docx alignment enum to string."""
    if alignment is None:
        return None
    align_map = {0: "left", 1: "center", 2: "right", 3: "justified"}
    if isinstance(alignment, int):
        return align_map.get(alignment)
    return str(alignment).lower().split(".")[-1] if alignment else None


def _extract_line_spacing(paragraph) -> Optional[float]:
    """Extract line spacing from paragraph format."""
    pf = paragraph.paragraph_format
    if pf.line_spacing is None:
        return None
    ls = float(pf.line_spacing)
    if pf.line_spacing_rule is not None:
        rule = str(pf.line_spacing_rule)
        if "MULTIPLE" in rule or "PROPORTIONAL" in rule:
            return ls
        elif "EXACTLY" in rule or "AT_LEAST" in rule:
            return round(ls / 12, 2)  # Convert pt to approximate multiplier
    return ls


def _is_heading(paragraph, style_map: dict) -> tuple[bool, Optional[int]]:
    """Detect if paragraph is a heading and its level."""
    style_name = paragraph.style.name if paragraph.style else ""
    style_name_lower = style_name.lower()

    # Check by style name
    if "heading" in style_name_lower:
        for i in range(1, 7):
            if f"heading {i}" == style_name_lower or style_name_lower == f"heading{i}":
                return True, i
        match = re.search(r"(\d)", style_name_lower)
        if match:
            return True, int(match.group(1))
        return True, 1

    # Heuristic: short bold text could be a heading
    text = paragraph.text.strip()
    if text and len(text) < 200:
        runs = paragraph.runs
        if runs and all(r.bold for r in runs if r.text.strip()):
            font_size = runs[0].font.size
            if font_size and font_size.pt >= 13:
                return True, 1
            elif font_size and font_size.pt >= 12:
                if len(text) < 100:
                    return True, 2

    return False, None


def _is_caption(paragraph) -> bool:
    """Detect if paragraph is a figure/table caption."""
    style_name = (paragraph.style.name or "").lower()
    if "caption" in style_name:
        return True
    text = paragraph.text.strip().lower()
    if text.startswith(("figure ", "fig. ", "fig ", "table ", "tbl ")):
        return True
    return False


def _is_reference(paragraph, in_references_section: bool) -> bool:
    """Detect if paragraph is a reference entry."""
    if in_references_section:
        text = paragraph.text.strip()
        if text and len(text) > 10:
            return True
    style_name = (paragraph.style.name or "").lower()
    if "bibliography" in style_name or "reference" in style_name:
        return True
    return False


def _estimate_page_numbers(paragraphs_count: int, page_count: int) -> list[int]:
    """Estimate page number for each paragraph based on even distribution."""
    if page_count <= 0:
        page_count = 1
    pages = []
    per_page = max(1, paragraphs_count // page_count)
    for i in range(paragraphs_count):
        pages.append(min((i // per_page) + 1, page_count))
    return pages


def _extract_page_setup(doc: Document) -> PageSetup:
    """Extract page setup from the document's section properties."""
    try:
        section = doc.sections[0] if doc.sections else None
        if section:
            width = section.page_width.cm if section.page_width else 21.0
            height = section.page_height.cm if section.page_height else 29.7
            orientation = "landscape" if width > height else "portrait"
            return PageSetup(
                width_cm=round(width, 2),
                height_cm=round(height, 2),
                margin_top_cm=round(section.top_margin.cm, 2) if section.top_margin else 2.54,
                margin_bottom_cm=round(section.bottom_margin.cm, 2) if section.bottom_margin else 2.54,
                margin_left_cm=round(section.left_margin.cm, 2) if section.left_margin else 2.54,
                margin_right_cm=round(section.right_margin.cm, 2) if section.right_margin else 2.54,
                orientation=orientation,
                paper_size=_detect_paper_size(width, height),
            )
    except Exception as e:
        logger.warning("Failed to extract page setup", error=str(e))

    return PageSetup(
        width_cm=21.0, height_cm=29.7,
        margin_top_cm=2.54, margin_bottom_cm=2.54,
        margin_left_cm=2.54, margin_right_cm=2.54,
        orientation="portrait", paper_size="A4",
    )


def _extract_header_footer(doc: Document) -> HeaderFooterContent:
    """Extract header/footer content and page number position."""
    has_header = False
    has_footer = False
    header_text = None
    footer_text = None
    page_number_position = None
    page_number_format = "arabic"

    try:
        for section in doc.sections:
            # Header
            header = section.header
            if header and not header.is_linked_to_previous:
                h_text = "\n".join(p.text for p in header.paragraphs if p.text.strip())
                if h_text.strip():
                    has_header = True
                    header_text = h_text.strip()
                # Check for page numbers in header
                for p in header.paragraphs:
                    for run in p.runs:
                        if run._element.xml and "fldChar" in run._element.xml:
                            page_number_position = "top_center"
                            break

            # Footer
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                f_text = "\n".join(p.text for p in footer.paragraphs if p.text.strip())
                if f_text.strip():
                    has_footer = True
                    footer_text = f_text.strip()
                # Check for page numbers in footer
                for p in footer.paragraphs:
                    xml_str = p._element.xml if hasattr(p._element, 'xml') else ""
                    if "PAGE" in xml_str or "fldChar" in xml_str or any(
                        c.isdigit() for c in p.text
                    ):
                        alignment = _get_alignment_str(p.paragraph_format.alignment)
                        if alignment == "center":
                            page_number_position = "bottom_center"
                        elif alignment == "right":
                            page_number_position = "bottom_right"
                        else:
                            page_number_position = "bottom_left"
                        break
    except Exception as e:
        logger.warning("Failed to extract header/footer", error=str(e))

    return HeaderFooterContent(
        has_header=has_header,
        has_footer=has_footer,
        header_text=header_text,
        footer_text=footer_text,
        page_number_position=page_number_position,
        page_number_format=page_number_format,
    )


def _detect_toc(doc: Document) -> bool:
    """Detect Table of Contents by looking for SDT elements or TOC headings."""
    try:
        body = doc.element.body
        # Look for structured document tags (SDT) with TOC
        for sdt in body.iter(qn("w:sdt")):
            sdt_pr = sdt.find(qn("w:sdtPr"))
            if sdt_pr is not None:
                doc_part = sdt_pr.find(qn("w:docPartGallery"))
                if doc_part is not None and "Table of Contents" in (doc_part.get(qn("w:val")) or ""):
                    return True

        # Heuristic: look for paragraphs with TOC styles
        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower()
            if style_name.startswith("toc") or "table of contents" in style_name:
                return True
            if para.text.strip().lower() in ("table of contents", "contents"):
                return True
    except Exception:
        pass
    return False


def parse_docx(file_bytes: bytes, filename: str) -> StructuredDocumentObject:
    """
    Main entry point. Takes raw .docx bytes, returns StructuredDocumentObject.
    Raises ValueError if the file is invalid.
    """
    # Step 1: Validate
    _validate_docx(file_bytes)

    # Step 2: Open document
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Failed to open .docx file: {str(e)}")

    # Step 3: Build style inheritance map
    style_map = _build_style_map(doc)

    # Step 4: Extract page setup
    page_setup = _extract_page_setup(doc)

    # Step 5: Extract header/footer
    header_footer = _extract_header_footer(doc)

    # Step 6: Detect TOC
    has_toc = _detect_toc(doc)

    # Step 7: Process all paragraphs
    paragraphs = []
    word_count = 0
    detected_sections = []
    has_references = False
    in_references_section = False

    # Estimate page count (rough: ~25 paragraphs per page for body text)
    total_paras = len(doc.paragraphs)
    estimated_pages = max(1, total_paras // 25)
    page_estimates = _estimate_page_numbers(total_paras, estimated_pages)

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        word_count += len(text.split()) if text else 0

        # Determine style
        style_id = para.style.style_id if para.style else None
        style_name = para.style.name if para.style else None

        # Check if heading
        is_heading_flag, heading_level = _is_heading(para, style_map)

        # Track sections
        if is_heading_flag and text:
            detected_sections.append(text)
            if re.match(r"(?i)(references?|bibliography|works cited)", text):
                has_references = True
                in_references_section = True
            elif is_heading_flag and heading_level == 1:
                in_references_section = False

        # Extract font properties with inheritance resolution
        font_family = None
        font_size_pt = None
        bold = None
        italic = None
        underline = None
        color = None

        # First: check inline run properties
        if para.runs:
            first_run = para.runs[0]
            if first_run.font.name:
                font_family = first_run.font.name
            if first_run.font.size:
                font_size_pt = first_run.font.size.pt
            bold = first_run.font.bold
            italic = first_run.font.italic
            underline = first_run.font.underline is not None and first_run.font.underline
            if first_run.font.color and first_run.font.color.rgb:
                color = str(first_run.font.color.rgb)

        # Second: resolve from style inheritance
        if font_family is None and style_id:
            font_family = _resolve_property(style_map, style_id, "font_family")
        if font_size_pt is None and style_id:
            resolved = _resolve_property(style_map, style_id, "font_size_pt")
            if resolved:
                font_size_pt = resolved
        if bold is None and style_id:
            bold = _resolve_property(style_map, style_id, "bold")
        if italic is None and style_id:
            italic = _resolve_property(style_map, style_id, "italic")

        # Extract paragraph formatting
        pf = para.paragraph_format
        alignment = _get_alignment_str(pf.alignment)
        if alignment is None and style_id:
            alignment = _resolve_property(style_map, style_id, "alignment")

        line_spacing = _extract_line_spacing(para)
        if line_spacing is None and style_id:
            line_spacing = _resolve_property(style_map, style_id, "line_spacing")

        space_before = pf.space_before.pt if pf.space_before else None
        space_after = pf.space_after.pt if pf.space_after else None

        indent_left = pf.left_indent.cm if pf.left_indent else None
        indent_right = pf.right_indent.cm if pf.right_indent else None
        indent_first = pf.first_line_indent.cm if pf.first_line_indent else None

        # Numbering
        numbering_id = None
        numbering_level = None
        is_list_item = False
        try:
            num_pr = para._element.find(qn("w:pPr"))
            if num_pr is not None:
                num_ref = num_pr.find(qn("w:numPr"))
                if num_ref is not None:
                    is_list_item = True
                    ilvl = num_ref.find(qn("w:ilvl"))
                    num_id = num_ref.find(qn("w:numId"))
                    if ilvl is not None:
                        numbering_level = int(ilvl.get(qn("w:val"), 0))
                    if num_id is not None:
                        numbering_id = int(num_id.get(qn("w:val"), 0))
        except Exception:
            pass

        page_num = page_estimates[idx] if idx < len(page_estimates) else estimated_pages

        paragraphs.append(ParagraphFormatting(
            paragraph_index=idx,
            page_number=page_num,
            text_preview=text[:100] if text else "",
            style_name=style_name,
            font_family=font_family,
            font_size_pt=font_size_pt,
            bold=bold,
            italic=italic,
            underline=underline,
            color=color,
            alignment=alignment,
            line_spacing=line_spacing,
            space_before_pt=space_before,
            space_after_pt=space_after,
            indent_left_cm=round(indent_left, 2) if indent_left else None,
            indent_right_cm=round(indent_right, 2) if indent_right else None,
            indent_first_line_cm=round(indent_first, 2) if indent_first else None,
            numbering_id=numbering_id,
            numbering_level=numbering_level,
            is_heading=is_heading_flag,
            heading_level=heading_level,
            is_caption=_is_caption(para),
            is_reference=_is_reference(para, in_references_section),
            is_list_item=is_list_item,
        ))

    # Store raw XML for correction engine
    raw_xml = None
    try:
        raw_xml = etree.tostring(doc.element.body, encoding="unicode")
    except Exception:
        pass

    return StructuredDocumentObject(
        filename=filename,
        page_count=estimated_pages,
        word_count=word_count,
        paragraph_count=len(paragraphs),
        page_setup=page_setup,
        header_footer=header_footer,
        paragraphs=paragraphs,
        has_table_of_contents=has_toc,
        has_references_section=has_references,
        detected_sections=detected_sections,
        raw_xml=raw_xml,
    )
